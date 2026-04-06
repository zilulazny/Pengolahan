def main():
    # Semua kode yang ada di dalam Jupyter dimasukkan ke sini
    import numpy as np
    import pygmt
    import pandas as pd
    import re
    import matplotlib.pyplot as plt
    import os
    import tkinter as tk
    import shutil  # Untuk menyalin file
    import locale
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.shared import Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION, WD_ORIENT
    from openpyxl import load_workbook
    from tkinter import filedialog
    from tkinter import messagebox
    from tkinter import ttk
    from tkcalendar import DateEntry 

    # Variabel global untuk menyimpan file yang dipilih
    selected_files = []

    # Fungsi untuk memilih file
    def select_files():
        global selected_files
        selected_files = filedialog.askopenfilenames(title="Pilih File", filetypes=(("Excel Files", "*.xlsx"), ("Semua File", "*.*")))
        if selected_files:
            messagebox.showinfo("File Terpilih", f"Jumlah file terpilih: {len(selected_files)}")

    # Fungsi untuk mengganti singkatan bulan dari Indonesia ke Inggris
    def replace_indonesian_months(date_str):
        ind_to_eng_months = {
            'Jan': 'Jan',
            'Feb': 'Feb',
            'Mar': 'Mar',
            'Apr': 'Apr',
            'Mei': 'May',
            'Jun': 'Jun',
            'Jul': 'Jul',
            'Agu': 'Aug',
            'Sep': 'Sep',
            'Okt': 'Oct',
            'Nov': 'Nov',
            'Des': 'Dec'
        }
        for ind, eng in ind_to_eng_months.items():
            date_str = date_str.replace(ind, eng)
        return date_str

    # Fungsi untuk memproses data
    def process_data():
        global selected_files
        
        if not selected_files:
            messagebox.showerror("Error", "Pilih file terlebih dahulu!")
            return
        
        start_date_str = start_date_entry.get()
        end_date_str = end_date_entry.get()
        
        start_date = datetime.strptime(start_date_str, '%d/%m/%Y')
        end_date = datetime.strptime(end_date_str, '%d/%m/%Y')
        
        column_name = 'Info Gempa'
        
        def parse_text(text):
            pattern = r"Mag:(?P<magnitude>\d+\.\d+) , (?P<date>\d{2}-\w{3}-\d{2}) (?P<time>\d{2}:\d{2}:\d{2}) WIB, Lok:(?P<latitude>\d+\.\d+) LS,(?P<longitude>\d+\.\d+) BT \((?P<description>.+?)\), Kedlmn:(?P<depth>\d+) Km"
            match = re.search(pattern, text)
            if match:
                result = match.groupdict()
                if 'LS' in text:
                    result['latitude'] = '-' + result['latitude']
                result['date'] = replace_indonesian_months(result['date'])
                return result
            return None
        
        combined_data = pd.DataFrame()
        
        for file in selected_files:
            df = pd.read_excel(file, header=1)
            parsed_data = df[column_name].dropna().apply(parse_text).apply(pd.Series)
            parsed_data = parsed_data.rename(columns={
                'date': 'Tanggal',
                'time': 'Waktu (WIB)',
                'latitude': 'Lintang',
                'longitude': 'Bujur',
                'magnitude': 'Magnitude',
                'depth': 'Kedalaman',
                'description': 'Keterangan'
            })
            
            parsed_data = parsed_data[['Tanggal', 'Waktu (WIB)', 'Lintang', 'Bujur', 'Kedalaman','Magnitude', 'Keterangan']]
            parsed_data['Dirasakan'] = ''
            parsed_data['Tanggal'] = pd.to_datetime(parsed_data['Tanggal'], format='%d-%b-%y')
            
            start_longitude = 113.21
            end_longitude = 117.31
            
            mask = (
                (parsed_data['Tanggal'] >= start_date) & 
                (parsed_data['Tanggal'] <= end_date) &
                (parsed_data['Bujur'].astype(float) >= start_longitude) & 
                (parsed_data['Bujur'].astype(float) <= end_longitude)
            )
            filtered_data = parsed_data[mask]
            
            combined_data = pd.concat([combined_data, filtered_data])
        
        # Simpan hasil gabungan ke file sementara
        combined_data.to_csv('dataspk.csv', index=False)
        
        # Baca data dari dataspk.csv
        hasil_data = pd.read_csv('dataspk.csv')

        # Baca data dari file Excel yang sudah ada
        existing_file = 'dirasakan.xlsx'
        if not os.path.exists(existing_file):
            messagebox.showerror("Error", f"File '{existing_file}' tidak ditemukan!")
            return

        existing_data = pd.read_excel(existing_file)

        # Pastikan nama kolom sama antara 'hasil_data' dan 'existing_data'
        existing_data = existing_data.rename(columns={
            'Tanggal': 'Tanggal',
            'Waktu(WIB)': 'Waktu (WIB)',
            'Magnitude': 'Magnitude',
            'Lintang': 'Lintang',
            'Bujur': 'Bujur',
            'Kedalaman': 'Kedalaman',
            'Keterangan': 'Keterangan',
            'Dirasakan': 'Dirasakan'
        })

        # Konversi kolom Tanggal ke datetime
        hasil_data['Tanggal'] = pd.to_datetime(hasil_data['Tanggal']).dt.date
        existing_data['Tanggal'] = pd.to_datetime(existing_data['Tanggal']).dt.date

        # Konversi kolom Waktu (WIB) menjadi string
        hasil_data['Waktu (WIB)'] = hasil_data['Waktu (WIB)'].astype(str)
        existing_data['Waktu (WIB)'] = existing_data['Waktu (WIB)'].astype(str)

        # Gabungkan data berdasarkan kunci unik
        key_columns = ['Tanggal', 'Waktu (WIB)', 'Magnitude', 'Lintang', 'Bujur', 'Kedalaman']
        merged_data = pd.merge(existing_data, hasil_data, on=key_columns, how='outer', suffixes=('', '_new'))

        # Prioritaskan data dari hasil_data (suffix '_new') jika ada data yang sama
        for col in hasil_data.columns:
            if col not in key_columns:
                merged_data[col] = merged_data[col + '_new'].combine_first(merged_data[col])
                merged_data.drop(columns=[col + '_new'], inplace=True)

        # Gabungkan kolom Tanggal dan Waktu (WIB) menjadi satu kolom Datetime untuk sort
        merged_data['Datetime'] = pd.to_datetime(merged_data['Tanggal'].astype(str) + ' ' + merged_data['Waktu (WIB)'])

        # Urutkan berdasarkan kolom Datetime
        merged_data = merged_data.sort_values(by='Datetime')

        # Hapus kolom Datetime setelah pengurutan
        merged_data.drop(columns=['Datetime'], inplace=True)

        # Tentukan nama file output berdasarkan tanggal awal dan akhir
        start_date_formatted = start_date.strftime('%d-%m-%Y')
        end_date_formatted = end_date.strftime('%d-%m-%Y')
        output_filename = f'{start_date_formatted}_{end_date_formatted}.csv'
        
        # Simpan file kedua di folder yang dipilih pengguna
        folder_tujuan = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan File")
        if not folder_tujuan:
            messagebox.showerror("Error", "Folder tujuan tidak dipilih!")
            return
        
        output_file = os.path.join(folder_tujuan, output_filename)
        merged_data.to_csv(output_file, index=False)

        # Salin file output ke 'inputpeta.csv' di direktori script
        try:
            shutil.copy(output_file, os.path.join(os.path.dirname(os.path.abspath(output_file)), 'D:/latsar/aktualisasi/coba/Untitled Folder/datalengkap.csv'))
        except Exception as e:
            messagebox.showerror("Error", f"Error saat menyalin file: {e}")
        
        messagebox.showinfo("Proses", f"Data berhasil diproses dan disimpan ke:\n1. {output_file}\n2. datalengkap.csv di direktori script.")

    # Fungsi untuk mencetak peta
    def print_map():
       #membuat file .dat
        df = pd.read_csv('datalengkap.csv')

        # Memilih kolom yang diinginkan
        selected_columns = df[['Bujur','Lintang','Kedalaman', 'Magnitude']]

        # Menulis data ke file DAT
        with open('inputpeta.dat', 'w') as file:
            for index, row in selected_columns.iterrows():
                line = '  '.join(row.astype(str))  # Menggabungkan semua elemen baris menjadi satu string dengan spasi sebagai pemisah
                file.write(line + '\n')
                
        # Membaca data dari hasil2.csv
        data = pd.read_csv('datalengkap.csv')

        # Buat figure
        fig = pygmt.Figure()

        # Buat color palette untuk kedalaman gempa
        pygmt.makecpt(cmap="geo", series=[-7000, 4000])
        pygmt.makecpt(cmap="red,yellow,green", series="0,60,300,1000", output="quakes.cpt")

        # Plot peta dasar dengan topografi
        fig.grdimage(
            grid="@earth_relief_15s",
            region=[113.21, 117.31, -12, -7],
            projection="M17c",
            shading=True,
            frame=["xa2g2", "ya2g2"])#, '+t"Bali"'])

        # Plot gempa yang tidak dirasakan (kolom 'Dirasakan' kosong)
        gempa_tidak_dirasakan = data[data['Dirasakan'].isna()]
        fig.plot(
            x=gempa_tidak_dirasakan['Bujur'],
            y=gempa_tidak_dirasakan['Lintang'],
            size=0.08*gempa_tidak_dirasakan['Magnitude'],
            style="cc",  # Lingkaran
            fill=gempa_tidak_dirasakan['Kedalaman'].astype(float),  # Warna berdasarkan kedalaman
            cmap="quakes.cpt",  # Gunakan colormap yang telah ditentukan
            pen="black"
        )

        # Plot gempa yang dirasakan (kolom 'Dirasakan' terisi)
        gempa_dirasakan = data[~data['Dirasakan'].isna()]
        fig.plot(
            x=gempa_dirasakan['Bujur'],
            y=gempa_dirasakan['Lintang'],
            size=0.08*gempa_dirasakan['Magnitude'],
            style="a0.5c",  # Simbol bintang
            fill="red",  # Warna bintang merah
            pen="1p,black"  # Garis tepi hitam dengan ketebalan 1 point
        )

        # Plot fault lines
        fig.plot(data="D:/latsar/aktualisasi/fault.gmt", pen="0.6,black")
        fig.plot(data="D:/latsar/aktualisasi/trench.gmt", pen="4,black")

        # Menambahkan colorbar
        fig.image('D:/latsar/aktualisasi/peta/mtangn.png', position='n0.79/0.83+w1.1i')
        fig.basemap(map_scale="n0.89/0.25+w80k+f0.1p+lKm")
        fig.colorbar(
            frame=["x+lKetinggian", "y"],  # Label untuk sumbu x dan y
            position="n0.87/0.3+w3c/0.5c")  # Posisi colorbar (JMR = tengah-kanan), lebar 0.5 cm, panjang 10 cm

        # Menambahkan legenda (diletakkan sebelum cross section)
        fig.plot(
            x=[115.8, 115.8, 117.31, 117.31],  # Koordinat X dari sudut-sudut persegi
            y=[-11, -12, -12, -11],  # Koordinat Y dari sudut-sudut persegi
            fill="white",  # Warna latar belakang
            close=True  # Menutup jalur untuk membuat persegi berisi warna 
        )
        fig.plot(x=[117.31, 115.8, 115.8], y=[-11, -11, -12], pen="1p,black")
        fig.legend(spec="D:/latsar/aktualisasi/coba/legend.gmt", position="n0.63/-0.015+w5c")

        with fig.inset(position="jBL+w5c/2.5c+o0.2c", box="+pblack"):
            fig.coast(
                region=[97, 140, -15, 7],
                shorelines='0.5p,black',
                projection="M5c",
                land="green",
                water="lightblue"
            )
            fig.plot(x=[117.31, 113.21, 113.21, 117.31, 117.31], y=[-7, -7, -12, -12, -7], pen="1p,red")

        # Membuat cross section
        fig.plot(x=[114.8, 115.4], y=[-11.5, -7.5], projection="M", pen=2)
        fig.text(x=114.9, y=-11.6, text="A", font="18,Helvetica")
        fig.text(x=115.55, y=-7.5, text="A1", font="18,Helvetica")
        pygmt.project(
            data="inputpeta.dat",
            unit=True,
            center=[114.8, -11.5],
            endpoint=[115.4, -7.5],
            width=[-200, 200],
            convention="pz",
            outfile="crsx8.dat"
        )

        fig.shift_origin(yshift=3, xshift=0.2)
        fig.plot(
            x=[113, 113, 114.18, 114.18],  # Koordinat X dari sudut-sudut persegi
            y=[-11.4, -12, -12, -11.4],  # Koordinat Y dari sudut-sudut persegi
            fill="white",  # Warna latar belakang
            close=True  # Menutup jalur untuk membuat persegi berisi warna 
        )

        # Buat peta cross section dengan basemap
        fig.basemap(
            projection="X4c/-2.5c",
            region=[0, 450, 0, 300],
            frame=['xafg+l"Distance (km)"','yafg+l"Depth (km)"', "wsEN"]
        )

        # Membuat garis cross section    
        fig.plot(data="crsx8.dat", projection="X", style="c0.15", pen=0.3, fill='red')
        fig.text(x=20, y=25, text="A", font="11,Helvetica")
        fig.text(x=370, y=25, text="A1", font="11,Helvetica")

        # Tentukan nama file peta berdasarkan tanggal awal dan akhir
        start_date_formatted = start_date_entry.get().replace('/', '-')  # Ganti format tanggal
        end_date_formatted = end_date_entry.get().replace('/', '-')
        output_map_filename = f"D:/latsar/aktualisasi/coba/Untitled Folder/peta/peta_seismisitas_{start_date_formatted}-{end_date_formatted}.jpg"

        # Menyimpan gambar
        fig.savefig(fname=output_map_filename)
        #fig.show()
        
        messagebox.showinfo("Cetak Peta", "Peta Berhasil di Cetak")

    # Fungsi untuk mencetak analisis xlsx
    def print_analysis():
        
        # Pastikan direktori output ada
        base_output_dir = r'D:\latsar\aktualisasi\coba\Untitled Folder\analisis'

        def read_csv_and_analyze(start_date, end_date):
            # Step 1: Read CSV file
            file_path = "datalengkap.csv"
            data = pd.read_csv(file_path)
        
            # Ensure 'Tanggal' column is in datetime format for grouping
            data['Tanggal'] = pd.to_datetime(data['Tanggal']).dt.date  # Menghapus jam, hanya menyisakan tanggal
            
            # Ambil tanggal awal dan akhir dari data
            start_date = data['Tanggal'].min()
            end_date = data['Tanggal'].max()
        
            # Drop the 'Magnitude Category' and 'Depth Category' columns
            output_data = data.drop(columns=['Magnitude Category', 'Depth Category'], errors='ignore')
        
            # Create magnitude and depth categories
            categories_magnitude = ['M<3', '3≤M<5', 'M≥5']
            data['Magnitude Category'] = pd.cut(data['Magnitude'], bins=[0, 3, 5, float('inf')], labels=categories_magnitude)
        
            categories_depth = ['D≤60 km', '60<D≤300 km', 'D>300 km']
            data['Depth Category'] = pd.cut(data['Kedalaman'], bins=[0, 60, 300, float('inf')], labels=categories_depth)
        
            # Group data by day and categories for analysis
            magnitude_summary = data.groupby([data['Tanggal'], 'Magnitude Category']).size().unstack(fill_value=0)
            depth_summary = data.groupby([data['Tanggal'], 'Depth Category']).size().unstack(fill_value=0)
        
            # Add totals row and columns for 'Gempa Dirasakan' and 'Gempa Merusak'
            magnitude_summary['Jumlah Total'] = magnitude_summary.sum(axis=1)
            magnitude_summary['Gempa Dirasakan'] = 0
            magnitude_summary['Gempa Merusak'] = 0
            magnitude_summary.loc['Jumlah Gempa'] = magnitude_summary.sum()
        
            depth_summary['Jumlah Total'] = depth_summary.sum(axis=1)
            depth_summary['Gempa Dirasakan'] = 0
            depth_summary['Gempa Merusak'] = 0
            depth_summary.loc['Jumlah Gempa'] = depth_summary.sum()
        
            # Mengisi kolom 'Gempa Dirasakan' jika kolom 'Dirasakan' terisi
            for index, row in data.iterrows():
                if pd.notna(row['Dirasakan']) and row['Dirasakan'].strip() != '':  # Cek apakah kolom 'Dirasakan' tidak kosong
                    # Tambahkan jumlah gempa dirasakan pada tanggal yang sesuai
                    if row['Tanggal'] in magnitude_summary.index:
                        magnitude_summary.at[row['Tanggal'], 'Gempa Dirasakan'] += 1
                    if row['Tanggal'] in depth_summary.index:
                        depth_summary.at[row['Tanggal'], 'Gempa Dirasakan'] += 1
        
            # Menambahkan jumlah total gempa dirasakan ke baris 'Jumlah Gempa'
            magnitude_summary.at['Jumlah Gempa', 'Gempa Dirasakan'] = magnitude_summary['Gempa Dirasakan'].sum()
            depth_summary.at['Jumlah Gempa', 'Gempa Dirasakan'] = depth_summary['Gempa Dirasakan'].sum()
        
            # Buat folder baru berdasarkan tanggal
            start_date_formatted = start_date_entry.get().replace('/', '-')  # Ganti format tanggal
            end_date_formatted = end_date_entry.get().replace('/', '-')
            folder_name = f"{start_date_formatted}-{end_date_formatted}"
            output_dir = os.path.join(base_output_dir, folder_name)
            os.makedirs(output_dir, exist_ok=True)

            # Save the analysis to Excel
            output_file = os.path.join(output_dir, f'{start_date_formatted}-{end_date_formatted}.xlsx')
            with pd.ExcelWriter(output_file, engine='xlsxwriter') as writer:
                # Save data to Excel
                output_data.to_excel(writer, sheet_name='Earthquake Data', index=False)
            
                # Save magnitude and depth summaries
                magnitude_summary.to_excel(writer, sheet_name='Magnitude Summary')
                depth_summary.to_excel(writer, sheet_name='Depth Summary')

                # Plotting the bar charts and pie charts
                plot_charts(magnitude_summary, depth_summary, writer, output_dir)

            # Open the saved Excel file and autofit columns
            autofit_columns(output_file, 'Earthquake Data')
            autofit_columns(output_file, 'Magnitude Summary')
            autofit_columns(output_file, 'Depth Summary')

            print(f'Analysis saved to {output_file}')

        def plot_charts(magnitude_summary, depth_summary, writer, output_dir):
            # Set the size for both bar and pie charts
            bar_chart_size = (16, 10)
            pie_chart_size = (5, 5)

            # Plot magnitude bar chart
            plt.figure(figsize=bar_chart_size)
            ax = magnitude_summary.drop(columns=['Jumlah Total', 'Gempa Dirasakan', 'Gempa Merusak']).drop(index='Jumlah Gempa').plot(kind='bar', stacked=True, color=['green', 'yellow', 'red'], edgecolor='black')
            ax.set_title('Berdasarkan Magnitudo', fontsize=18)
            ax.set_xlabel("Tanggal", fontsize=14)
            ax.set_ylabel("Frekuensi Kejadian", fontsize=14)
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right', fontsize=12)
            ax.legend(title='Kategori Magnitudo', fontsize=7, title_fontsize='9', loc='upper right')
            ax.grid(axis='y', linestyle='--', alpha=0.7)
        
            # Manual label placement at the bottom inside each bar
            for p in ax.patches:
                height = p.get_height()
                if height > 0:
                    ax.annotate(f'{int(height)}', 
                                (p.get_x() + p.get_width() / 2, p.get_y() + 0.1 * height),
                                ha='center', va='bottom', fontsize=10, color='black')
        
            plt.tight_layout(pad=2)
            plt.savefig(os.path.join(output_dir, 'diagbat_mag.png'))
            plt.close()

            # Insert the magnitude bar chart into Excel
            worksheet_mag = writer.sheets['Magnitude Summary']
            worksheet_mag.insert_image('A10', os.path.join(output_dir, 'diagbat_mag.png'), {'x_scale': 1.4, 'y_scale': 1.6})

            # Plot magnitude pie chart
            plt.figure(figsize=pie_chart_size)
            magnitude_total = magnitude_summary.loc['Jumlah Gempa', ['M<3', '3≤M<5', 'M≥5']]
            magnitude_total.plot(kind='pie', autopct=lambda pct: f'{pct:.1f}%\n({int(pct/100*sum(magnitude_total))} kejadian)', startangle=90, colors=['green', 'yellow', 'red'])
            plt.title('Distribusi Magnitudo', fontsize=11)
            plt.tight_layout()
            plt.savefig(os.path.join(output_dir, 'diaglingk_mag.png'))
            plt.close()

            # Insert the magnitude pie chart into Excel
            worksheet_mag.insert_image('O10', os.path.join(output_dir, 'diaglingk_mag.png'), {'x_scale': 1.2, 'y_scale': 1.2})

            # Plot depth bar chart
            plt.figure(figsize=bar_chart_size)
            ax = depth_summary.drop(columns=['Jumlah Total', 'Gempa Dirasakan', 'Gempa Merusak']).drop(index='Jumlah Gempa').plot(kind='bar', stacked=True, color=['red', 'yellow', 'green'], edgecolor='black')
            ax.set_title('Berdasarkan Kedalaman', fontsize=18)
            ax.set_xlabel("Tanggal", fontsize=14)
            ax.set_ylabel("Frekuensi Kejadian", fontsize=14)
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right', fontsize=12)
            ax.legend(title='Kategori Kedalaman', fontsize=7, title_fontsize='9', loc='upper right')
            ax.grid(axis='y', linestyle='--', alpha=0.7)
        
            # Manual label placement at the bottom inside each bar
            for p in ax.patches:
                height = p.get_height()
                if height > 0:
                    ax.annotate(f'{int(height)}', 
                                (p.get_x() + p.get_width() / 2, p.get_y() + 0.1 * height),
                                ha='center', va='bottom', fontsize=10, color='black')
        
            plt.tight_layout(pad=2)
            plt.savefig(os.path.join(output_dir, 'diagbat_depth.png'))
            plt.close()

            # Insert the depth bar chart into Excel
            worksheet_depth = writer.sheets['Depth Summary']
            worksheet_depth.insert_image('A10', os.path.join(output_dir, 'diagbat_depth.png'), {'x_scale': 1.4, 'y_scale': 1.6})

            # Plot depth pie chart
            plt.figure(figsize=pie_chart_size)
            depth_total = depth_summary.loc['Jumlah Gempa', ['D≤60 km', '60<D≤300 km', 'D>300 km']]
            depth_total.plot(kind='pie', autopct=lambda pct: f'{pct:.1f}%\n({int(pct/100*sum(depth_total))} kejadian)', startangle=90, colors=['red', 'yellow', 'green'])
            plt.title('Distribusi Kedalaman', fontsize=11)
            plt.tight_layout()
            plt.savefig(os.path.join(output_dir, 'diaglingk_depth.png'))
            plt.close()

            # Insert the depth pie chart into Excel
            worksheet_depth.insert_image('O10', os.path.join(output_dir, 'diaglingk_depth.png'), {'x_scale': 1.2, 'y_scale': 1.2})

        def autofit_columns(file_path, sheet_name):
            """
            Menyesuaikan lebar kolom di Excel berdasarkan panjang konten.
            """
            workbook = load_workbook(file_path)  # Membuka workbook yang telah disimpan
            worksheet = workbook[sheet_name]  # Mendapatkan worksheet berdasarkan nama sheet

            for col in worksheet.columns:
                max_length = 0
                column = col[0].column_letter  # Mendapatkan huruf kolom, misalnya 'A', 'B', dst.
                for cell in col:
                    try:
                        # Mencari panjang maksimum dari setiap kolom
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)  # Menambah padding
                worksheet.column_dimensions[column].width = adjusted_width  # Mengatur lebar kolom

            workbook.save(file_path)  # Menyimpan workbook setelah penyesuaian
            workbook.close()  # Menutup workbook untuk memastikan semua perubahan disimpan

        # Panggil fungsi dengan tanggal awal dan akhir yang sesuai
        # Misalnya:
        read_csv_and_analyze(start_date_entry, end_date_entry)

        messagebox.showinfo("Cetak Analisis", "Analisis Excel Berhasil di Cetak")


    # Fungsi untuk mencetak laporan Word
    def print_report():
        # Set locale to Indonesian for date formatting
        locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')

        # Fungsi untuk memilih file
        def select_file():
            root = tk.Tk()
            root.withdraw()  # Menyembunyikan jendela utama
            file_path = filedialog.askopenfilename(title="Pilih file laporan_gempa.xlsx", filetypes=[("Excel files", "*.xlsx")])
            return file_path

        # Memilih file Excel
        file_path = select_file()
        if not file_path:
            print("Tidak ada file yang dipilih.")
        else:
            # Load the Excel file
            earthquake_data = pd.read_excel(file_path)

            # Pastikan kolom 'Tanggal' dalam format datetime
            earthquake_data['Tanggal'] = pd.to_datetime(earthquake_data['Tanggal'])

            # Ambil tanggal awal dan akhir
            start_date = earthquake_data['Tanggal'].min()
            end_date = earthquake_data['Tanggal'].max()

            # Format tanggal untuk digunakan dalam nama folder

            folder_name = f"{start_date.strftime('%d-%m-%Y')}-{end_date.strftime('%d-%m-%Y')}"
            print(folder_name)


            # Path to images
            logo_path = "LogoBMKG.png"  # Ganti dengan path logo BMKG
            magnitude_bar_chart_path = f"D:/latsar/aktualisasi/coba/Untitled Folder/analisis/{folder_name}/diagbat_mag.png"
            magnitude_pie_chart_path = f"D:/latsar/aktualisasi/coba/Untitled Folder/analisis/{folder_name}/diaglingk_mag.png"
            depth_bar_chart_path = f"D:/latsar/aktualisasi/coba/Untitled Folder/analisis/{folder_name}/diagbat_depth.png"
            depth_pie_chart_path = f"D:/latsar/aktualisasi/coba/Untitled Folder/analisis/{folder_name}/diaglingk_depth.png"
            seismic_map_path = f"D:/latsar/aktualisasi/coba/Untitled Folder/peta/peta_seismisitas_{folder_name}.jpg"
            print(seismic_map_path)

            earthquake_data = pd.read_excel(file_path)

            # Pastikan kolom 'Tanggal' dalam format datetime
            earthquake_data['Tanggal'] = pd.to_datetime(earthquake_data['Tanggal'])

            # Calculate the required values
            total_earthquakes = len(earthquake_data)
            min_magnitude = earthquake_data['Magnitude'].min()
            max_magnitude = earthquake_data['Magnitude'].max()

            # Date range
            start_date = earthquake_data['Tanggal'].min()
            end_date = earthquake_data['Tanggal'].max()
            #date1 = earthquake_data['Tanggal']

            # Format tanggal dengan nama bulan dalam bahasa Indonesia
            start_date_formatted = start_date.strftime('%d %B %Y')  # contoh format: 01 Januari 2024
            end_date_formatted = end_date.strftime('%d %B %Y')      # contoh format: 05 September 2024
            #date_formatted=date1.strftime('%d %B %Y')
            #print(date_formatted)

            # Menentukan tanggal awal bulan untuk memulai perhitungan minggu
            start_of_month = start_date.replace(day=1)

            # Hitung nomor minggu relatif terhadap awal bulan
            week_number = (start_date - start_of_month).days // 7 + 1

            # Format bulan dalam bahasa Indonesia
            bulan_formatted = start_date.strftime('%B')

            # Menghitung jumlah kejadian gempabumi total
            total_earthquakes = len(earthquake_data)

            # Menghitung jumlah gempa dengan lat < -9 dan kedalaman < 60 km
            shallow_earthquake_count = earthquake_data[(earthquake_data['Lintang'] < -9) & 
                                                       (earthquake_data['Kedalaman'] < 60)].shape[0]

            # Menghitung jumlah gempa di selatan Pulau Jawa, Bali, dan Lombok
            southern_earthquake_count = total_earthquakes - shallow_earthquake_count

            # Magnitude categories
            magnitude_below_3 = earthquake_data[earthquake_data['Magnitude'] < 3]
            magnitude_3_to_5 = earthquake_data[(earthquake_data['Magnitude'] >= 3) & (earthquake_data['Magnitude'] <= 5)]
            magnitude_above_5 = earthquake_data[earthquake_data['Magnitude'] > 5]

            count_below_3 = len(magnitude_below_3)
            count_3_to_5 = len(magnitude_3_to_5)
            count_above_5 = len(magnitude_above_5)

            percentage_below_3 = int((count_below_3 / total_earthquakes) * 100)
            percentage_3_to_5 = int((count_3_to_5 / total_earthquakes) * 100)
            percentage_above_5 = int((count_above_5 / total_earthquakes) * 100)

            # Magnitude categories
            depth_below_60 = earthquake_data[earthquake_data['Kedalaman'] < 60]
            depth_60_to_300 = earthquake_data[(earthquake_data['Kedalaman'] >= 60) & (earthquake_data['Kedalaman'] <= 300)]
            depth_above_300 = earthquake_data[earthquake_data['Kedalaman'] > 300]

            count_depth_below_60 = len(depth_below_60)
            count_depth_60_to_300 = len(depth_60_to_300)
            count_depth_above_300 = len(depth_above_300)

            percentage_depth_below_60 = int((count_depth_below_60 / total_earthquakes) * 100)
            percentage_depth_60_to_300 = int((count_depth_60_to_300 / total_earthquakes) * 100)
            percentage_depth_above_300 = int((count_depth_above_300 / total_earthquakes) * 100)

            # Determine dominant depth category
            dominant_category = ''
            if percentage_depth_below_60 >= percentage_depth_60_to_300 and percentage_depth_below_60 >= percentage_depth_above_300:
                dominant_category = f'dangkal (<60 km) sebesar {percentage_depth_below_60}%'
            elif percentage_depth_60_to_300 >= percentage_depth_below_60 and percentage_depth_60_to_300 >= percentage_depth_above_300:
                dominant_category = f'kedalaman menengah (60 - 300 km) sebesar {percentage_depth_60_to_300}%'
            else:
                dominant_category = f'kedalaman dalam (>300 km) sebesar {percentage_depth_above_300}%'

            # Create a new Word document
            document = Document()

            # Ukuran kertas F4 (21.0 cm x 33.0 cm)
            f4_width = Cm(21.0)
            f4_height = Cm(33.0)

            # Mengatur ukuran kertas F4 untuk setiap section di dokumen
            for section in document.sections:
                # Mengatur orientasi potrait untuk F4
                section.orientation = WD_ORIENT.PORTRAIT  # Bisa diubah ke WD_ORIENT.LANDSCAPE jika ingin landscape
                section.page_width = f4_width
                section.page_height = f4_height

                # Mengatur margin halaman (optional, bisa disesuaikan)
                section.top_margin = Cm(2.5)     # Marginn atas
                section.bottom_margin = Cm(2.5)  # Marginn bawah
                section.left_margin = Cm(2.5)    # Marginn kiri
                section.right_margin = Cm(2.5)   # Marginn kanan

            # Add Logo at the top
            logo_paragraph = document.add_paragraph()
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(1))  # Menambahkan logo dengan lebar 1.5 inches
            logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Title
            title = document.add_paragraph()
            title_run = title.add_run('STASIUN GEOFISIKA DENPASAR')
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.all_caps = True  # Membuat semua huruf kapital
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Subtitle
            subtitle = document.add_paragraph()
            subtitle_run = subtitle.add_run('SEISMISITAS WILAYAH BALI DAN SEKITARNYA')
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.all_caps = True  # Membuat semua huruf kapital
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Date Period
            period = document.add_paragraph()
            period_run = period.add_run(f'PERIODE {start_date_formatted} – {end_date_formatted}')
            period_run.font.size = Pt(12)
            period_run.font.all_caps = True  # Membuat semua huruf kapital
            period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add space
            document.add_paragraph('')

            # Teks Deskriptif Magnitudo dengan perhitungan minggu bulanan
            magnitude_descriptive_text = (
                f"Berdasarkan data Stasiun Geofisika Denpasar selama minggu ke-{week_number} bulan "
                f"{bulan_formatted} {start_date.year}, "
                f"di daerah Bali dan sekitarnya telah terjadi {total_earthquakes} kejadian gempabumi dengan magnitudo bervariasi "
                f"mulai dari M {min_magnitude} sampai M {max_magnitude}. Kejadian gempabumi dengan magnitudo M<3 sejumlah "
                f"{count_below_3} kejadian atau {percentage_below_3}% dari total kejadian gempabumi. Sedangkan untuk M 3 – 5 "
                f"sejumlah {count_3_to_5} kejadian atau {percentage_3_to_5}% dari total kejadian gempabumi "
            )

            # Conditional text for magnitude above 5
            if count_above_5 > 0:
                magnitude_descriptive_text += (
                    f"dan kejadian gempa M>5 sejumlah {count_above_5} kejadian atau {percentage_above_5}% dari total kejadian gempabumi."
                )
            else:
                magnitude_descriptive_text += "dan tidak terdapat kejadian gempabumi dengan magnitudo M>5."

            # Menambahkan paragraf dengan teks deskriptif magnitudo
            magnitude_paragraph = document.add_paragraph(magnitude_descriptive_text)

            # Mengatur alignment paragraf menjadi justify
            magnitude_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add Magnitude Bar Chart Image
            try:
                document.add_paragraph().add_run().add_picture(magnitude_bar_chart_path, width=Inches(4))
                document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except FileNotFoundError:
                document.add_paragraph('DIAGRAM BATANG MAGNITUDO TIDAK DITEMUKAN').bold = True

            # Add Magnitude Pie Chart Image
            try:
                document.add_paragraph().add_run().add_picture(magnitude_pie_chart_path, width=Inches(3))
                document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except FileNotFoundError:
                document.add_paragraph('DIAGRAM LINGKARAN MAGNITUDO TIDAK DITEMUKAN').bold = True

            # Teks Keterangan Gambar Magnitudo
            document.add_paragraph('Gambar 1. Histogram gempabumi harian berdasarkan magnitudo (atas) dan diagram gempabumi berdasarkan magnitudo (bawah)').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Determine dominant depth category
            dominant_category = ''
            if percentage_depth_below_60 >= percentage_depth_60_to_300 and percentage_depth_below_60 >= percentage_depth_above_300:
                dominant_category = f'dangkal (<60 km) sebesar {percentage_depth_below_60}%'
            elif percentage_depth_60_to_300 >= percentage_depth_below_60 and percentage_depth_60_to_300 >= percentage_depth_above_300:
                dominant_category = f'kedalaman menengah (60 - 300 km) sebesar {percentage_depth_60_to_300}%'
            else:
                dominant_category = f'kedalaman dalam (>300 km) sebesar {percentage_depth_above_300}%'

            # Add space
            document.add_paragraph('')

            # Teks Deskriptif Kedalaman
            depth_descriptive_text = (
                f"Berdasarkan kedalaman, gempabumi yang mendominasi adalah kejadian gempabumi {dominant_category}  dari total kejadian gempabumi. "
                f"Jumlah gempabumi dangkal sebanyak {count_depth_below_60} kejadian gempabumi, "
                f"terdapat sebanyak {count_depth_60_to_300} kejadian gempabumi dengan kedalaman menengah 60 km - 300 km atau {percentage_depth_60_to_300}% dari total kejadian gempabumi, "
            )

            # Conditional text for depth above 300 km
            if count_depth_above_300 > 0:
                depth_descriptive_text += (
                    f"dan kejadian gempa dalam sejumlah {count_depth_above_300} kejadian atau {percentage_depth_above_300}% dari total kejadian gempabumi."
                )
            else:
                depth_descriptive_text += "dan tidak terdapat kejadian gempabumi dengan kedalaman dalam >300 km."

            depth_paragraph = document.add_paragraph(depth_descriptive_text)
            depth_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add Depth Bar Chart Image
            try:
                document.add_paragraph().add_run().add_picture(depth_bar_chart_path, width=Inches(4))
                document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except FileNotFoundError:
                document.add_paragraph('DIAGRAM BATANG KEDALAMAN TIDAK DITEMUKAN').bold = True

            # Add Depth Pie Chart Image
            try:
                document.add_paragraph().add_run().add_picture(depth_pie_chart_path, width=Inches(3))
                document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except FileNotFoundError:
                document.add_paragraph('DIAGRAM LINGKARAN KEDALAMAN TIDAK DITEMUKAN').bold = True

            # Teks Keterangan Gambar Kedalaman (akan diisi sesuai format)
            document.add_paragraph('Gambar 2.  Histogram gempabumi harian berdasarkan kedalaman (atas) dan diagram gempabumi berdasarkan kedalaman (bawah)').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Seismic Map Image
            try:
                document.add_paragraph().add_run().add_picture(seismic_map_path, width=Inches(5))
                document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except FileNotFoundError:
                document.add_paragraph('PETA SEISMISITAS TIDAK DITEMUKAN').bold = True

            # Teks Keterangan Gambar Peta Seismisitas
            peta_seismisitas_text = (
                f"Gambar 3. Peta seismisitas Bali dan sekitarnya periode {start_date_formatted} - {end_date_formatted}"
            )
            peta_seismisitas_paragraph = document.add_paragraph(peta_seismisitas_text)
            peta_seismisitas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add space
            document.add_paragraph('')

            # Teks analisis yang diminta
            analisis_text1 = (
                f"Selama satu minggu terakhir, gempabumi yang terjadi di Bali dan sekitarnya dikelompokan berdasarkan sumbernya (Gambar 3):\n"
            )    

            # Menambahkan paragraf analisis ke dalam dokumen
            analisis_paragraph1 = document.add_paragraph(analisis_text1)
            analisis_paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            analisis_text2 = (
                f"1. Sebanyak {southern_earthquake_count} kejadian gempabumi terjadi di selatan Pulau Jawa, Bali, dan Lombok "
                f"yang diperkirakan berasosiasi dengan subduksi Indo Australia-Eurasia.\n"
                f"2. Sebanyak {shallow_earthquake_count} kejadian gempabumi terjadi menyebar di wilayah Pulau Bali, Lombok dan sekitarnya "
                f"diperkirakan berkaitan dengan sesar di belakang busur Flores atau Flores Back Arc Thrust dan aktivitas sesar aktif."
            )

            # Menambahkan paragraf analisis ke dalam dokumen
            analisis_paragraph2 = document.add_paragraph(analisis_text2)
            analisis_paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Menghitung jumlah kejadian gempa yang dirasakan
            earthquakes_felt = earthquake_data[earthquake_data['Dirasakan'].notna()]
            total_felt = len(earthquakes_felt)

            # Menyusun narasi pembuka
            if total_felt > 0:
                felt_intro = (
                    f"Selama periode ini terdapat {total_felt} kejadian gempabumi yang dirasakan oleh masyarakat "
                    f"di wilayah Bali dan sekitarnya."
                )
            else:
                felt_intro = (
                    f"Selama periode ini tidak terdapat kejadian gempabumi yang dirasakan oleh masyarakat "
                    f"di wilayah Bali dan sekitarnya."
                )

            # Menambahkan narasi pembuka ke dokumen
            felt_intro_paragraph = document.add_paragraph(felt_intro)
            felt_intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Menyusun detail setiap kejadian gempa yang dirasakan dalam format list dengan nomor urut
            if total_felt > 0:
                # Inisialisasi counter untuk nomor urut
                counter = 1

                for index, row in earthquakes_felt.iterrows():
                    magnitude = row['Magnitude']
                    date = row['Tanggal'].strftime('%d %B %Y')  # Format tanggal DD/MM/YYYY
                    time = row['Waktu (WIB)']  # Asumsikan kolom waktu sudah dalam format string yang benar
                    location = row['Dirasakan']
                    description = row['Keterangan']
                    depth = row['Kedalaman']

                    # Format detail gempa yang dirasakan dengan nomor urut
                    felt_detail = (
                        f"{counter}. Gempabumi dengan magnitudo {magnitude} dirasakan di wilayah {location} yang terjadi pada tanggal {date} "
                        f"pukul {time} WIB berlokasi di {description} pada kedalaman {depth} km."
                    )

                    # Menambahkan detail gempa ke dalam dokumen sebagai paragraf terpisah
                    felt_detail_paragraph = document.add_paragraph(felt_detail)
                    felt_detail_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    # Meningkatkan counter untuk nomor urut berikutnya
                    counter += 1
            else:
                # Jika tidak ada gempa yang dirasakan, tambahkan keterangan tambahan
                no_felt_detail = document.add_paragraph("Tidak ada rincian gempabumi yang dirasakan selama periode ini.")
                no_felt_detail.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            # Tambahkan section baru dengan orientasi landscape
            new_section = document.add_section(WD_SECTION.NEW_PAGE)
            new_section.orientation = WD_ORIENT.LANDSCAPE

            # Atur ukuran kertas agar landscape sesuai
            new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

            # Menentukan margin untuk section baru agar tabel berada di tengah secara vertikal
            section_margin = Cm(2)  # Sesuaikan margin jika perlu
            new_section.top_margin = section_margin
            new_section.bottom_margin = section_margin
            new_section.left_margin = section_margin
            new_section.right_margin = section_margin

            # Menambahkan paragraf baru untuk membungkus tabel
            table_paragraph = document.add_paragraph()
            table_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Mengatur alignment paragraf ke tengah

            # Membuat teks keterangan tabel gempabumi
            tabel_keterangan_text = (
                f"Tabel 1. Data gempabumi di Bali dan sekitarnya tanggal {start_date_formatted} – {end_date_formatted}"
            )

            # Menambahkan teks keterangan tabel ke dalam dokumen
            tabel_keterangan_paragraph = document.add_paragraph(tabel_keterangan_text)
            tabel_keterangan_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Mengatur alignment menjadi tengah

            # Menambahkan tabel langsung ke dalam paragraf
            table = document.add_table(rows=1, cols=len(earthquake_data.columns))
            table.style = 'Table Grid'

            # Set table alignment to center
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths to match the text content
            column_widths = {
                "NO.": 1.5,  # NO
                "Tanggal": 2.5,  # TANGGAL
                "Waktu (WIB)": 2.5,  # WAKTU (WIB)
                "Lintang": 2.0,  # LATITUDE
                "Bujur": 2.0,  # LONGITUDE
                "Kedalaman": 2.0,  # Depth
                "Magnitude": 1.5,  # Magnitude
                "Keterangan": 5.0,  # Keterangan
                "Dirasakan": 2.5  # Dirasakan
            }

            # Add header row
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(earthquake_data.columns):
                hdr_cells[i].text = column_name
                hdr_cells[i].width = Cm(column_widths.get(column_name, 2.0))  # Set width based on column

            # Add data rows, memformat kolom 'Tanggal' tanpa jam
            for index, row in earthquake_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    # Jika kolom adalah 'Tanggal', hanya tampilkan tanggal tanpa jam
                    if earthquake_data.columns[i] == 'Tanggal':
                        row_cells[i].text = value.strftime('%d/%m/%Y')  # Format tanggal menjadi DD/MM/YYYY
                    # Jika kolom adalah 'Dirasakan' dan nilai adalah NaN, ganti dengan teks yang sesuai
                    elif earthquake_data.columns[i] == 'Dirasakan':
                        # Ganti NaN dengan string kosong atau teks lain
                        row_cells[i].text = '' if pd.isna(value) else str(value)
                    else:
                        row_cells[i].text = str(value)

                    # Set the width of each data cell, sesuai dengan kolom
                    row_cells[i].width = Cm(column_widths.get(earthquake_data.columns[i], 2.0))

            # Mengatur tabel agar berada di tengah secara horizontal
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Save the document
            document_name = f'D:/latsar/aktualisasi/coba/Untitled Folder/laporan/Laporan Gempabumi {folder_name}.docx'
            document.save(document_name)
            print(f'DOKUMEN BERHASIL DISIMPAN SEBAGAI {document_name}')

        messagebox.showinfo("Cetak Laporan", "Mencetak laporan ke file Word...")

    # Fungsi untuk mencetak deskripsi
    def print_description():
        # Set locale to Indonesian for date formatting
        locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')

        # Fungsi untuk memilih file
        def select_file():
            root = tk.Tk()
            root.withdraw()  # Menyembunyikan jendela utama
            file_path = filedialog.askopenfilename(title="Pilih file laporan_gempa.xlsx", filetypes=[("Excel files", "*.xlsx")])
            return file_path

        # Memilih file Excel
        file_path = select_file()
        if not file_path:
            print("Tidak ada file yang dipilih.")
        else:
            # Load the Excel file
            earthquake_data = pd.read_excel(file_path)

            # Pastikan kolom 'Tanggal' dalam format datetime
            earthquake_data['Tanggal'] = pd.to_datetime(earthquake_data['Tanggal'])

            # Ambil tanggal awal dan akhir
            start_date = earthquake_data['Tanggal'].min()
            end_date = earthquake_data['Tanggal'].max()

            # Format tanggal untuk digunakan dalam nama folder

            folder_name = f"{start_date.strftime('%d-%m-%Y')}-{end_date.strftime('%d-%m-%Y')}"
            print(folder_name)

            earthquake_data = pd.read_excel(file_path)

            # Pastikan kolom 'Tanggal' dalam format datetime
            earthquake_data['Tanggal'] = pd.to_datetime(earthquake_data['Tanggal'])

            # Calculate the required values
            total_earthquakes = len(earthquake_data)
            min_magnitude = earthquake_data['Magnitude'].min()
            max_magnitude = earthquake_data['Magnitude'].max()

            # Date range
            start_date = earthquake_data['Tanggal'].min()
            end_date = earthquake_data['Tanggal'].max()
            #date1 = earthquake_data['Tanggal']

            # Format tanggal dengan nama bulan dalam bahasa Indonesia
            start_date_formatted = start_date.strftime('%d %B %Y')  # contoh format: 01 Januari 2024
            end_date_formatted = end_date.strftime('%d %B %Y')      # contoh format: 05 September 2024
            #date_formatted=date1.strftime('%d %B %Y')
            #print(date_formatted)

            # Menentukan tanggal awal bulan untuk memulai perhitungan minggu
            start_of_month = start_date.replace(day=1)

            # Hitung nomor minggu relatif terhadap awal bulan
            week_number = (start_date - start_of_month).days // 7 + 1

            # Format bulan dalam bahasa Indonesia
            bulan_formatted = start_date.strftime('%B')

            # Menghitung jumlah kejadian gempabumi total
            total_earthquakes = len(earthquake_data)

            # Menghitung jumlah gempa dengan lat < -9 dan kedalaman < 60 km
            shallow_earthquake_count = earthquake_data[(earthquake_data['Lintang'] < -9) & 
                                                       (earthquake_data['Kedalaman'] < 60)].shape[0]

            # Menghitung jumlah gempa di selatan Pulau Jawa, Bali, dan Lombok
            southern_earthquake_count = total_earthquakes - shallow_earthquake_count

            # Magnitude categories
            magnitude_below_3 = earthquake_data[earthquake_data['Magnitude'] < 3]
            magnitude_3_to_5 = earthquake_data[(earthquake_data['Magnitude'] >= 3) & (earthquake_data['Magnitude'] <= 5)]
            magnitude_above_5 = earthquake_data[earthquake_data['Magnitude'] > 5]

            count_below_3 = len(magnitude_below_3)
            count_3_to_5 = len(magnitude_3_to_5)
            count_above_5 = len(magnitude_above_5)

            percentage_below_3 = int((count_below_3 / total_earthquakes) * 100)
            percentage_3_to_5 = int((count_3_to_5 / total_earthquakes) * 100)
            percentage_above_5 = int((count_above_5 / total_earthquakes) * 100)

            # Magnitude categories
            depth_below_60 = earthquake_data[earthquake_data['Kedalaman'] < 60]
            depth_60_to_300 = earthquake_data[(earthquake_data['Kedalaman'] >= 60) & (earthquake_data['Kedalaman'] <= 300)]
            depth_above_300 = earthquake_data[earthquake_data['Kedalaman'] > 300]

            count_depth_below_60 = len(depth_below_60)
            count_depth_60_to_300 = len(depth_60_to_300)
            count_depth_above_300 = len(depth_above_300)

            percentage_depth_below_60 = int((count_depth_below_60 / total_earthquakes) * 100)
            percentage_depth_60_to_300 = int((count_depth_60_to_300 / total_earthquakes) * 100)
            percentage_depth_above_300 = int((count_depth_above_300 / total_earthquakes) * 100)

            # Determine dominant depth category
            dominant_category = ''
            if percentage_depth_below_60 >= percentage_depth_60_to_300 and percentage_depth_below_60 >= percentage_depth_above_300:
                dominant_category = f'dangkal (<60 km) sebesar {percentage_depth_below_60}%'
            elif percentage_depth_60_to_300 >= percentage_depth_below_60 and percentage_depth_60_to_300 >= percentage_depth_above_300:
                dominant_category = f'kedalaman menengah (60 - 300 km) sebesar {percentage_depth_60_to_300}%'
            else:
                dominant_category = f'kedalaman dalam (>300 km) sebesar {percentage_depth_above_300}%'

            # Create a new Word document
            document = Document()

            # Ukuran kertas F4 (21.0 cm x 33.0 cm)
            f4_width = Cm(21.0)
            f4_height = Cm(33.0)

            # Mengatur ukuran kertas F4 untuk setiap section di dokumen
            for section in document.sections:
                # Mengatur orientasi potrait untuk F4
                section.orientation = WD_ORIENT.PORTRAIT  # Bisa diubah ke WD_ORIENT.LANDSCAPE jika ingin landscape
                section.page_width = f4_width
                section.page_height = f4_height

                # Mengatur margin halaman (optional, bisa disesuaikan)
                section.top_margin = Cm(2.5)     # Marginn atas
                section.bottom_margin = Cm(2.5)  # Marginn bawah
                section.left_margin = Cm(2.5)    # Marginn kiri
                section.right_margin = Cm(2.5)   # Marginn kanan

            # Add Title
            title = document.add_paragraph()
            title_run = title.add_run('STASIUN GEOFISIKA DENPASAR')
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.all_caps = True  # Membuat semua huruf kapital
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Subtitle
            subtitle = document.add_paragraph()
            subtitle_run = subtitle.add_run('SEISMISITAS WILAYAH BALI DAN SEKITARNYA')
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.all_caps = True  # Membuat semua huruf kapital
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add Date Period
            period = document.add_paragraph()
            period_run = period.add_run(f'PERIODE {start_date_formatted} – {end_date_formatted}')
            period_run.font.size = Pt(12)
            period_run.font.all_caps = True  # Membuat semua huruf kapital
            period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add space
            document.add_paragraph('')

            # Teks Deskriptif Magnitudo dengan perhitungan minggu bulanan
            magnitude_descriptive_text = (
                f"Berdasarkan data Stasiun Geofisika Denpasar selama minggu ke-{week_number} bulan "
                f"{bulan_formatted} {start_date.year}, "
                f"di daerah Bali dan sekitarnya telah terjadi {total_earthquakes} kejadian gempabumi dengan magnitudo bervariasi "
                f"mulai dari M {min_magnitude} sampai M {max_magnitude}. Kejadian gempabumi dengan magnitudo M<3 sejumlah "
                f"{count_below_3} kejadian atau {percentage_below_3}% dari total kejadian gempabumi. Sedangkan untuk M 3 – 5 "
                f"sejumlah {count_3_to_5} kejadian atau {percentage_3_to_5}% dari total kejadian gempabumi "
            )

            # Conditional text for magnitude above 5
            if count_above_5 > 0:
                magnitude_descriptive_text += (
                    f"dan kejadian gempa M>5 sejumlah {count_above_5} kejadian atau {percentage_above_5}% dari total kejadian gempabumi."
                )
            else:
                magnitude_descriptive_text += "dan tidak terdapat kejadian gempabumi dengan magnitudo M>5."

            # Menambahkan paragraf dengan teks deskriptif magnitudo
            magnitude_paragraph = document.add_paragraph(magnitude_descriptive_text)

            # Mengatur alignment paragraf menjadi justify
            magnitude_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Determine dominant depth category
            dominant_category = ''
            if percentage_depth_below_60 >= percentage_depth_60_to_300 and percentage_depth_below_60 >= percentage_depth_above_300:
                dominant_category = f'dangkal (<60 km) sebesar {percentage_depth_below_60}%'
            elif percentage_depth_60_to_300 >= percentage_depth_below_60 and percentage_depth_60_to_300 >= percentage_depth_above_300:
                dominant_category = f'kedalaman menengah (60 - 300 km) sebesar {percentage_depth_60_to_300}%'
            else:
                dominant_category = f'kedalaman dalam (>300 km) sebesar {percentage_depth_above_300}%'

            # Add space
            document.add_paragraph('')

            # Teks Deskriptif Kedalaman
            depth_descriptive_text = (
                f"Berdasarkan kedalaman, gempabumi yang mendominasi adalah kejadian gempabumi {dominant_category}  dari total kejadian gempabumi. "
                f"Jumlah gempabumi dangkal sebanyak {count_depth_below_60} kejadian gempabumi, "
                f"terdapat sebanyak {count_depth_60_to_300} kejadian gempabumi dengan kedalaman menengah 60 km - 300 km atau {percentage_depth_60_to_300}% dari total kejadian gempabumi, "
            )

            # Conditional text for depth above 300 km
            if count_depth_above_300 > 0:
                depth_descriptive_text += (
                    f"dan kejadian gempa dalam sejumlah {count_depth_above_300} kejadian atau {percentage_depth_above_300}% dari total kejadian gempabumi."
                )
            else:
                depth_descriptive_text += "dan tidak terdapat kejadian gempabumi dengan kedalaman dalam >300 km."

            depth_paragraph = document.add_paragraph(depth_descriptive_text)
            depth_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Menghitung jumlah kejadian gempa yang dirasakan
            earthquakes_felt = earthquake_data[earthquake_data['Dirasakan'].notna()]
            total_felt = len(earthquakes_felt)

            # Menyusun narasi pembuka
            if total_felt > 0:
                felt_intro = (
                    f"Selama periode ini terdapat {total_felt} kejadian gempabumi yang dirasakan oleh masyarakat "
                    f"di wilayah Bali dan sekitarnya."
                )
            else:
                felt_intro = (
                    f"Selama periode ini tidak terdapat kejadian gempabumi yang dirasakan oleh masyarakat "
                    f"di wilayah Bali dan sekitarnya."
                )

            # Menambahkan narasi pembuka ke dokumen
            felt_intro_paragraph = document.add_paragraph(felt_intro)
            felt_intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Menyusun detail setiap kejadian gempa yang dirasakan dalam format list dengan nomor urut
            if total_felt > 0:
                # Inisialisasi counter untuk nomor urut
                counter = 1

                for index, row in earthquakes_felt.iterrows():
                    magnitude = row['Magnitude']
                    date = row['Tanggal'].strftime('%d %B %Y')  # Format tanggal DD/MM/YYYY
                    time = row['Waktu (WIB)']  # Asumsikan kolom waktu sudah dalam format string yang benar
                    location = row['Dirasakan']
                    description = row['Keterangan']
                    depth = row['Kedalaman']

                    # Format detail gempa yang dirasakan dengan nomor urut
                    felt_detail = (
                        f"{counter}. Gempabumi dengan magnitudo {magnitude} dirasakan di wilayah {location} yang terjadi pada tanggal {date} "
                        f"pukul {time} WIB berlokasi di {description} pada kedalaman {depth} km."
                    )

                    # Menambahkan detail gempa ke dalam dokumen sebagai paragraf terpisah
                    felt_detail_paragraph = document.add_paragraph(felt_detail)
                    felt_detail_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    # Meningkatkan counter untuk nomor urut berikutnya
                    counter += 1
            else:
                # Jika tidak ada gempa yang dirasakan, tambahkan keterangan tambahan
                no_felt_detail = document.add_paragraph("Tidak ada rincian gempabumi yang dirasakan selama periode ini.")
                no_felt_detail.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



            # Save the document
            document_name = f'D:/latsar/aktualisasi/coba/Untitled Folder/deskripsi.docx'
            document.save(document_name)
            print(f'DOKUMEN BERHASIL DISIMPAN SEBAGAI {document_name}')
        messagebox.showinfo("Cetak Deskripsi", "Mencetak deskripsi...")

    # Membuat jendela utama
    root = tk.Tk()
    root.title("Aplikasi GUI Sederhana")
    root.geometry("400x400")

    # Tombol untuk memilih file
    select_file_button = ttk.Button(root, text="Pilih File", command=select_files)
    select_file_button.pack(pady=10)

    # Pilihan tanggal awal dan akhir
    ttk.Label(root, text="Pilih Tanggal Awal:").pack(pady=5)
    start_date_entry = DateEntry(root, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
    start_date_entry.pack(pady=5)

    ttk.Label(root, text="Pilih Tanggal Akhir:").pack(pady=5)
    end_date_entry = DateEntry(root, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
    end_date_entry.pack(pady=5)

    # Tombol proses
    process_button = ttk.Button(root, text="Proses", command=process_data)
    process_button.pack(pady=10)

    # Tombol cetak peta
    print_map_button = ttk.Button(root, text="Cetak Peta", command=print_map)
    print_map_button.pack(pady=5)

    # Tombol cetak analisis xlsx
    print_analysis_button = ttk.Button(root, text="Cetak Analisis xlsx", command=print_analysis)
    print_analysis_button.pack(pady=5)

    # Tombol cetak laporan Word
    print_report_button = ttk.Button(root, text="Cetak Laporan Word", command=print_report)
    print_report_button.pack(pady=5)

    # Tombol cetak deskripsi
    print_description_button = ttk.Button(root, text="Cetak Deskripsi", command=print_description)
    print_description_button.pack(pady=5)

    # Menjalankan aplikasi
    root.mainloop()

    # Fungsi dan class bisa tetap di luar main()

    # Mulai aplikasi GUI
    root = tk.Tk()
    root.title("Aplikasi GUI Sederhana")
    root.geometry("400x400")

    # Tombol untuk memilih file
    select_file_button = ttk.Button(root, text="Pilih File", command=select_files)
    select_file_button.pack(pady=10)

    # Pilihan tanggal awal dan akhir
    ttk.Label(root, text="Pilih Tanggal Awal:").pack(pady=5)
    start_date_entry = DateEntry(root, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
    start_date_entry.pack(pady=5)

    ttk.Label(root, text="Pilih Tanggal Akhir:").pack(pady=5)
    end_date_entry = DateEntry(root, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
    end_date_entry.pack(pady=5)

    # Tombol proses
    process_button = ttk.Button(root, text="Proses", command=process_data)
    process_button.pack(pady=10)

    # Tombol cetak peta
    print_map_button = ttk.Button(root, text="Cetak Peta", command=print_map)
    print_map_button.pack(pady=5)

    # Tombol cetak analisis xlsx
    print_analysis_button = ttk.Button(root, text="Cetak Analisis xlsx", command=print_analysis)
    print_analysis_button.pack(pady=5)

    # Tombol cetak laporan Word
    print_report_button = ttk.Button(root, text="Cetak Laporan Word", command=print_report)
    print_report_button.pack(pady=5)

    # Tombol cetak deskripsi
    print_description_button = ttk.Button(root, text="Cetak Deskripsi", command=print_description)
    print_description_button.pack(pady=5)

    # Menjalankan aplikasi
    root.mainloop()

if __name__ == "__main__":
    main()
